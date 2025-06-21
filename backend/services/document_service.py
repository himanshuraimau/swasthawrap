import io
import logging
from typing import Dict, Any, Optional
from pathlib import Path
import aiofiles

# PDF processing
try:
    import PyPDF2
    import pdfplumber
except ImportError:
    PyPDF2 = None
    pdfplumber = None

# Word document processing
try:
    from docx import Document
except ImportError:
    Document = None

logger = logging.getLogger(__name__)


class DocumentService:
    def __init__(self):
        self.supported_formats = {
            '.pdf': self.extract_pdf_text,
            '.docx': self.extract_docx_text,
            '.doc': self.extract_docx_text,
            '.txt': self.extract_txt_text
        }

    async def extract_text_from_file(self, file_path: str) -> Dict[str, Any]:
        """Extract text from various document formats"""
        try:
            file_extension = Path(file_path).suffix.lower()
            
            if file_extension not in self.supported_formats:
                return {
                    "text": "",
                    "error": f"Unsupported file format: {file_extension}",
                    "success": False
                }

            # Read file content
            async with aiofiles.open(file_path, 'rb') as file:
                content = await file.read()

            # Extract text based on file type
            extractor = self.supported_formats[file_extension]
            text = await extractor(content)

            return {
                "text": text,
                "word_count": len(text.split()),
                "char_count": len(text),
                "success": True,
                "file_type": file_extension
            }

        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {e}")
            return {
                "text": "",
                "error": str(e),
                "success": False
            }

    async def extract_pdf_text(self, content: bytes) -> str:
        """Extract text from PDF using multiple methods"""
        text = ""
        
        try:
            # Method 1: Try pdfplumber first (better for complex layouts)
            if pdfplumber:
                try:
                    with pdfplumber.open(io.BytesIO(content)) as pdf:
                        for page in pdf.pages:
                            page_text = page.extract_text()
                            if page_text:
                                text += page_text + "\n"
                    
                    if text.strip():
                        return text.strip()
                except Exception as e:
                    logger.warning(f"pdfplumber extraction failed: {e}")

            # Method 2: Fallback to PyPDF2
            if PyPDF2 and not text.strip():
                try:
                    pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
                    for page in pdf_reader.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                except Exception as e:
                    logger.warning(f"PyPDF2 extraction failed: {e}")

            return text.strip()

        except Exception as e:
            logger.error(f"PDF text extraction failed: {e}")
            return ""

    async def extract_docx_text(self, content: bytes) -> str:
        """Extract text from Word documents"""
        try:
            if not Document:
                raise ImportError("python-docx not available")

            doc = Document(io.BytesIO(content))
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"

            return text.strip()

        except Exception as e:
            logger.error(f"DOCX text extraction failed: {e}")
            return ""

    async def extract_txt_text(self, content: bytes) -> str:
        """Extract text from plain text files"""
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'iso-8859-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    return content.decode(encoding)
                except UnicodeDecodeError:
                    continue
            
            # If all encodings fail, use utf-8 with error handling
            return content.decode('utf-8', errors='replace')

        except Exception as e:
            logger.error(f"Text file extraction failed: {e}")
            return ""

    def chunk_text(self, text: str, max_chunk_size: int = 4000) -> list[str]:
        """Split text into chunks for processing"""
        if len(text) <= max_chunk_size:
            return [text]

        chunks = []
        words = text.split()
        current_chunk = []
        current_size = 0

        for word in words:
            word_size = len(word) + 1  # +1 for space
            if current_size + word_size > max_chunk_size and current_chunk:
                chunks.append(' '.join(current_chunk))
                current_chunk = [word]
                current_size = word_size
            else:
                current_chunk.append(word)
                current_size += word_size

        if current_chunk:
            chunks.append(' '.join(current_chunk))

        return chunks

    async def process_document_for_chat(self, file_path: str) -> Dict[str, Any]:
        """Process document and prepare it for chat integration"""
        try:
            # Extract text
            extraction_result = await self.extract_text_from_file(file_path)
            
            if not extraction_result["success"]:
                return extraction_result

            text = extraction_result["text"]
            
            # Create summary for better context
            summary = self.create_document_summary(text)
            
            # Chunk the text for processing
            chunks = self.chunk_text(text)

            return {
                "success": True,
                "text": text,
                "summary": summary,
                "chunks": chunks,
                "chunk_count": len(chunks),
                "word_count": extraction_result["word_count"],
                "char_count": extraction_result["char_count"],
                "file_type": extraction_result["file_type"]
            }

        except Exception as e:
            logger.error(f"Error processing document for chat: {e}")
            return {
                "success": False,
                "error": str(e)
            }

    def create_document_summary(self, text: str, max_length: int = 500) -> str:
        """Create a brief summary of the document"""
        if len(text) <= max_length:
            return text

        # Simple extractive summary - take first few sentences
        sentences = text.split('. ')
        summary = ""
        
        for sentence in sentences:
            if len(summary + sentence) <= max_length:
                summary += sentence + ". "
            else:
                break

        return summary.strip()


# Global instance
document_service = DocumentService()
